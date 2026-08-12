#!/usr/bin/env python3
"""Build the visual, Word-friendly CourtEdge ProGear presenter guide."""

from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "docs" / "demo-guide-assets"
OUTPUT = ROOT / "docs" / "CourtEdge-ProGear-Team-Demo-Guide.docx"

DEMO_URL = "https://progear-sales-aiagent.vercel.app"
VIDEO_URL = "https://drive.google.com/file/d/1N2XwGVgZXg2yHugp21hUTQoaBjaemxFw/view?usp=drive_link"

INK = "172033"
MUTED = "536078"
BROWN = "4A2C1A"
ORANGE = "E65F1B"
ORANGE_PALE = "FFF3EA"
PURPLE = "6437E8"
PURPLE_PALE = "F3EFFF"
BLUE = "2457D6"
BLUE_PALE = "EEF4FF"
GREEN = "147A4B"
GREEN_PALE = "EDF9F2"
RED = "B93A2B"
RED_PALE = "FFF0EE"
GOLD_PALE = "FFF8DB"
LINE = "D8DEE9"
PALE = "F7F8FB"
WHITE = "FFFFFF"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    element = tc_pr.find(qn("w:shd"))
    if element is None:
        element = OxmlElement("w:shd")
        tc_pr.append(element)
    element.set(qn("w:fill"), fill)


def cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_widths(table, widths: list[float]) -> None:
    """Set both the Word table grid and cell widths for stable rendering."""
    for index, width in enumerate(widths):
        table.columns[index].width = Inches(width)
        for cell in table.columns[index].cells:
            cell.width = Inches(width)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def set_cell_text(cell, text: str, *, bold=False, color=INK, size=9, align=None) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, label: str, url: str, *, color=BLUE) -> None:
    relation = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color_element = OxmlElement("w:color")
    color_element.set(qn("w:val"), color)
    properties.append(color_element)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_page(section, *, landscape=False) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(4)
    for name, size, color in (
        ("Title", 30, INK),
        ("Heading 1", 20, BROWN),
        ("Heading 2", 13, ORANGE),
        ("Heading 3", 10.5, PURPLE),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(4)
    caption = styles["Caption"]
    caption.font.name = "Arial"
    caption.font.size = Pt(7.5)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Arial"
    run.font.size = Pt(7.5)
    run.font.color.rgb = rgb(MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def configure_header_footer(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run("COURTEDGE PROGEAR  /  PRESENTER GUIDE")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(7.5)
    run.font.color.rgb = rgb(BROWN)
    footer = section.footer
    footer.is_linked_to_previous = False
    add_page_number(footer.paragraphs[0])


def add_new_section(document: Document, *, landscape=False) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    set_page(section, landscape=landscape)
    configure_header_footer(section)


def add_page_title(document: Document, number: str, title: str, subtitle: str) -> None:
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    section = document.sections[-1]
    available = (section.page_width - section.left_margin - section.right_margin) / 914400
    set_table_widths(table, [0.58, available - 0.58])
    number_cell, title_cell = table.rows[0].cells
    shade(number_cell, ORANGE)
    shade(title_cell, WHITE)
    cell_margins(number_cell, top=80, start=60, bottom=80, end=60)
    cell_margins(title_cell, top=20, start=130, bottom=20, end=40)
    set_cell_text(number_cell, number, bold=True, color=WHITE, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph = title_cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = rgb(BROWN)
    detail = title_cell.add_paragraph()
    detail.paragraph_format.space_after = Pt(0)
    run = detail.add_run(subtitle)
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = rgb(MUTED)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_label(document: Document, text: str, *, color=PURPLE) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(color)


def add_picture(document: Document, filename: str, width: float, caption: str) -> None:
    path = ASSETS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    caption_paragraph = document.add_paragraph(style="Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(4)
    caption_paragraph.add_run(caption)


def add_callout(document: Document, heading: str, text: str, *, fill=BLUE_PALE, accent=BLUE) -> None:
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    section = document.sections[-1]
    available = (section.page_width - section.left_margin - section.right_margin) / 914400
    set_table_widths(table, [1.18, available - 1.18])
    left, right = table.rows[0].cells
    shade(left, accent)
    shade(right, fill)
    cell_margins(left, top=100, start=90, bottom=100, end=90)
    cell_margins(right, top=100, start=130, bottom=100, end=130)
    set_cell_text(left, heading.upper(), bold=True, color=WHITE, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(right, text, color=INK, size=9)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_simple_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_table_widths(table, widths)
    header = table.rows[0]
    repeat_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        shade(cell, BROWN)
        cell_margins(cell, top=90, bottom=90)
        set_cell_text(cell, text, bold=True, color=WHITE, size=8.5)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for index, text in enumerate(values):
            cell = row.cells[index]
            shade(cell, WHITE if row_index % 2 == 0 else PALE)
            cell_margins(cell, top=80, bottom=80)
            set_cell_text(cell, text, size=8.5)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_steps(document: Document, steps: list[str], *, compact=False) -> None:
    table = document.add_table(rows=0, cols=2)
    table.autofit = False
    for index, text in enumerate(steps, start=1):
        row = table.add_row()
        prevent_row_split(row)
        number, detail = row.cells
        shade(number, ORANGE if index == 1 else BROWN)
        shade(detail, ORANGE_PALE if index == 1 else PALE)
        amount = 55 if compact else 75
        cell_margins(number, top=amount, start=35, bottom=amount, end=35)
        cell_margins(detail, top=amount, start=100, bottom=amount, end=100)
        set_cell_text(number, str(index), bold=True, color=WHITE, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(detail, text, size=8.5 if compact else 9)
    section = document.sections[-1]
    available = (section.page_width - section.left_margin - section.right_margin) / 914400
    set_table_widths(table, [0.42, available - 0.42])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_two_column_points(document: Document, points: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=0, cols=4)
    table.autofit = False
    pairs = [points[index:index + 2] for index in range(0, len(points), 2)]
    for pair in pairs:
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        for index in range(2):
            label_cell = cells[index * 2]
            text_cell = cells[index * 2 + 1]
            if index < len(pair):
                label, text = pair[index]
                shade(label_cell, BROWN)
                shade(text_cell, PALE)
                set_cell_text(label_cell, label, bold=True, color=WHITE, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_text(text_cell, text, size=8)
            else:
                shade(label_cell, WHITE)
                shade(text_cell, WHITE)
            for cell in (label_cell, text_cell):
                cell_margins(cell, top=70, start=75, bottom=70, end=75)
    section = document.sections[-1]
    available = (section.page_width - section.left_margin - section.right_margin) / 914400
    half = available / 2
    set_table_widths(table, [1.02, half - 1.02, 1.02, half - 1.02])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_cover(document: Document) -> None:
    rule = document.add_table(rows=1, cols=1)
    shade(rule.cell(0, 0), ORANGE)
    cell_margins(rule.cell(0, 0), top=18, bottom=18)
    rule.cell(0, 0).paragraphs[0].add_run(" ")

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(3)
    run = kicker.add_run("COURTEDGE PROGEAR")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = rgb(ORANGE)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Customer Demo Guide")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(28)
    run.font.color.rgb = rgb(INK)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(7)
    run = subtitle.add_run("A visual presenter runbook for Okta-governed AI agent access")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = rgb(MUTED)

    add_picture(document, "01-sign-in.png", 6.75, "Start here: the custom CourtEdge ProGear agent sign-in experience.")

    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    set_table_widths(table, [1.25, 5.85])
    values = (("LIVE DEMO", DEMO_URL), ("UPDATE VIDEO", VIDEO_URL))
    for index, (label, value) in enumerate(values):
        left, right = table.rows[index].cells
        shade(left, BROWN)
        shade(right, WHITE if index == 0 else GOLD_PALE)
        cell_margins(left, top=100, bottom=100)
        cell_margins(right, top=100, bottom=100)
        set_cell_text(left, label, bold=True, color=WHITE, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        paragraph = right.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        add_hyperlink(paragraph, "Open the live demo" if index == 0 else "Watch the recorded walkthrough", value)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run(f"Internal presenter guide  •  Updated {date.today().strftime('%B %-d, %Y')}  •  Keep credentials separate")
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(7.5)
    run.font.color.rgb = rgb(MUTED)


def add_demo_map(document: Document) -> None:
    document.add_page_break()
    add_page_title(document, "01", "Know the story before you click", "Two personas and one clear approval role")
    add_callout(
        document,
        "THE STORY",
        "Sarah can read but never write. With FGA on, Mike can add 1–600 units and a 601+ request requires a Level 2 VP.",
        fill=ORANGE_PALE,
        accent=ORANGE,
    )
    add_label(document, "Personas")
    add_simple_table(
        document,
        ["PERSONA", "OKTA PROFILE", "DEMO OUTCOME"],
        [
            ["Sarah Sales", "Clearance 0 • Manager False", "Reads Inventory. Every write is denied with manager guidance."],
            ["Mike Manager", "Clearance 1 • Manager True", "Writes 1–600 units. A 601+ request requires a VP."],
            ["Existing VP approver", "Clearance 2 • Manager True", "Approves Mike's 601+ request and may write directly."],
        ],
        [1.3, 2.0, 3.75],
    )
    add_label(document, "Core demo — about five minutes", color=ORANGE)
    add_steps(
        document,
        [
            "Sign in as Sarah and run the inventory read prompt.",
            "Ask Sarah to add 50 basketballs; show the clear Sales denial.",
            "Open Mike in the intended demo tab; a fresh tab starts with FGA off.",
            "Ask Mike to add 50 basketballs; show the successful write.",
            "In simple mode, ask Mike for a large positive write; show coarse Okta access succeeds.",
            "Open Architecture, then Request sequence, and walk left to right.",
            "Enable FGA and route Mike's 601+ request to AIAgentOwners.",
        ],
        compact=True,
    )
    add_callout(
        document,
        "DELEGATION",
        "On vacation is separate from role. When True, the agent stops before ID-JAG and cannot act for that employee.",
        fill=PURPLE_PALE,
        accent=PURPLE,
    )


def add_sarah_page(document: Document) -> None:
    document.add_page_break()
    add_page_title(document, "02", "Sarah: read succeeds, write stops", "Authentication identifies Sarah; her Okta role limits the action")
    add_label(document, "First, demonstrate the read", color=GREEN)
    add_picture(document, "03-sarah-read.png", 6.62, "Sarah asks for the basketball inventory. The read succeeds.")
    add_label(document, "Then, demonstrate the write", color=RED)
    add_picture(document, "04-sarah-write-denied.png", 6.62, "Sarah asks to add 50 units. Inventory does not change.")
    add_callout(
        document,
        "SAY THIS",
        "Sarah is Sales, Clearance 0. She can read Inventory, but the agent cannot turn her identity into write authority. She must contact her manager.",
        fill=RED_PALE,
        accent=RED,
    )


def add_mike_page(document: Document) -> None:
    document.add_page_break()
    add_page_title(document, "03", "Mike: coarse access first, then FGA", "The quantity boundary applies only after FGA is enabled")
    add_label(document, "1–600 units: execute", color=GREEN)
    add_picture(document, "06-mike-write-allowed.png", 6.62, "Mike adds 50 units. The response confirms the previous and new totals.")
    add_label(document, "FGA on — 601+ units: owner approval", color=PURPLE)
    add_picture(document, "07-mike-601-simple-denied.png", 6.62, "With FGA enabled, Mike's 601-unit request routes to AIAgentOwners.")
    add_callout(
        document,
        "SAY THIS",
        "Mike is a Manager, Clearance 1. Coarse Okta grants inventory:write without a quantity tier. The optional FGA demo adds the 600/601 boundary and AI Agent Owner approval route.",
        fill=PURPLE_PALE,
        accent=PURPLE,
    )


def add_architecture_page(document: Document) -> None:
    add_new_section(document, landscape=True)
    add_page_title(document, "04", "Architecture: preserve both identities", "The employee asks; the governed Workload Principal acts; each resource stays in control")
    add_picture(document, "08-architecture.png", 8.35, "Architecture view: employee, ProGear Agent, Okta, Resource AS, protected resources, audit trail, and kill switch.")
    add_two_column_points(
        document,
        [
            ("EMPLOYEE", "The signed-in person remains the subject of the request."),
            ("AGENT", "A first-class Workload Principal with an independent lifecycle."),
            ("CONTROL", "Okta preserves delegation; each resource receives scoped access."),
            ("AUDIT / KILL", "Decisions stay attributable; deactivation stops new exchanges."),
        ],
    )


def add_sequence_page(document: Document) -> None:
    document.add_page_break()
    add_page_title(document, "05", "Request sequence: follow the delegated proof", "ID-JAG bridges employee context and agent identity without merging them")
    add_picture(document, "09-sequence.png", 9.35, "Successful request sequence. Read the numbered messages from top to bottom.")
    add_two_column_points(
        document,
        [
            ("1–2", "Employee asks; the agent presents workload identity and user context."),
            ("3", "Okta issues an Identity Assertion Grant for this delegation."),
            ("4–5", "The agent exchanges it for a resource- and scope-specific token."),
            ("6–7", "Inventory returns the result and the chain remains auditable."),
        ],
    )
    add_callout(
        document,
        "BLOCKED PATH",
        "If delegation is stopped—such as On vacation=True—there is no ID-JAG and no scoped resource token.",
        fill=RED_PALE,
        accent=RED,
    )


def add_fga_controls_page(document: Document) -> None:
    add_new_section(document, landscape=False)
    add_page_title(document, "06", "Optional FGA demo: turn on the advanced path", "FGA is intentionally off until the presenter selects Simulate FGA")
    add_picture(document, "10-fga-controls.png", 5.92, "Open FGA as Mike, select Simulate FGA, and confirm Manager / Clearance 1 / On vacation False. Mike can also preview VP in this tab.")
    add_picture(document, "12-fga-prompts.png", 5.92, "The chat home now shows one prompt for each policy boundary: Read, 1–600, and 601+.")
    add_callout(
        document,
        "RESET RULE",
        "FGA is isolated to the current browser tab and survives refresh or sign-out. A new tab starts with FGA off. It does not modify Okta profile attributes.",
        fill=BLUE_PALE,
        accent=BLUE,
    )


def add_hitl_page(document: Document) -> None:
    document.add_page_break()
    add_page_title(document, "07", "Human In The Loop: one approval boundary", "Okta controls delegation; FGA controls the Inventory action; OIG records the owner decision")
    add_picture(document, "11-fga-policy.png", 6.25, "The policy combines live Okta attributes with role, action, and quantity.")
    add_simple_table(
        document,
        ["REQUEST", "RESULT"],
        [
            ["Sarah • any write", "Deny. Contact a Manager. Never create an approval request."],
            ["Mike • 1–600 units", "Execute directly."],
            ["Mike • 601+ units", "Create one AIAgentOwners request; do not change Inventory while pending."],
            ["Mike • VP preview", "Execute any quantity directly in this isolated demo session."],
        ],
        [2.0, 5.05],
    )


def add_approval_page(document: Document) -> None:
    document.add_page_break()
    add_page_title(document, "08", "Complete the owner approval and reset", "An AI Agent Owner approves in Okta; the backend verifies live group membership")
    add_picture(document, "13-oig-vp-approval.png", 6.72, "The Okta Access Requests task shows Mike as requester and presents Approve / Deny actions.")
    add_steps(
        document,
        [
            "As Mike with FGA enabled, select Add 601 basketballs to inventory.",
            "Copy the Okta request ID shown by ProGear.",
            "Open Okta Access Requests → Inbox → Open as Johnathan or another AI Agent Owner in a separate browser profile.",
            "Open Mike's request and select Approve.",
            "Return to ProGear. The backend verifies the approver's live AIAgentOwners membership and executes the change.",
        ],
        compact=True,
    )
    add_callout(
        document,
        "VACATION",
        "To show the delegation safeguard, set On vacation=True and issue a protected request. The agent stops before ID-JAG. Restore it to False before ending.",
        fill=PURPLE_PALE,
        accent=PURPLE,
    )
    add_label(document, "End-of-demo checklist", color=ORANGE)
    add_simple_table(
        document,
        ["RESET", "CONFIRM"],
        [
            ["Okta attributes", "Vacation False; Sarah 0, Mike 1; Manager False / True."],
            ["OIG", "Resolve or deny test requests that should not remain open."],
            ["Session", "Close the tab when finished; a new tab starts with FGA off."],
        ],
        [1.5, 5.55],
    )
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    run = paragraph.add_run("Recorded update: ")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    add_hyperlink(paragraph, "Watch the Google Drive walkthrough", VIDEO_URL)


def build() -> None:
    document = Document()
    document.core_properties.title = "CourtEdge ProGear Customer Demo Guide"
    document.core_properties.subject = "Visual step-by-step presenter guide"
    document.core_properties.author = "Johnathan Campos"
    document.core_properties.keywords = "Okta, AI Agent Governance, ID-JAG, FGA, OIG, ProGear"
    configure_styles(document)
    set_page(document.sections[0], landscape=False)
    document.sections[0].different_first_page_header_footer = True
    configure_header_footer(document.sections[0])

    add_cover(document)
    add_demo_map(document)
    add_sarah_page(document)
    add_mike_page(document)
    add_architecture_page(document)
    add_sequence_page(document)
    add_fga_controls_page(document)
    add_hitl_page(document)
    add_approval_page(document)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    try:
        build()
    except Exception as exc:
        print(f"Failed to build demo guide: {exc}", file=sys.stderr)
        raise
